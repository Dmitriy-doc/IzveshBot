from aiogram import Bot, Dispatcher, types
from aiogram.types import ReplyKeyboardMarkup, KeyboardButton
from aiogram.utils import executor
from aiogram.contrib.fsm_storage.memory import MemoryStorage
from aiogram.dispatcher import FSMContext
from aiogram.dispatcher.filters.state import State, StatesGroup

from docx import Document
from datetime import datetime
import os

# Токен берем из переменной окружения (например, на Render)
API_TOKEN = os.getenv("7518865505:AAEdCzkLa10pGA6N4uRyuy2CTDAQP0w-IOQ")

bot = Bot(token=API_TOKEN, parse_mode="HTML")
storage = MemoryStorage()
dp = Dispatcher(bot, storage=storage)

class Form(StatesGroup):
    doctor_name = State()
    patient_name = State()
    diagnosis   = State()
    lab_results = State()

start_keyboard = ReplyKeyboardMarkup(resize_keyboard=True)
start_keyboard.add(KeyboardButton("📝 Подать извещение"))

@dp.message_handler(commands=['start'])
async def cmd_start(message: types.Message):
    await message.answer("Нажмите кнопку '📝 Подать извещение', чтобы начать.", reply_markup=start_keyboard)

@dp.message_handler(lambda msg: msg.text == "📝 Подать извещение", state='*')
async def start_form(message: types.Message):
    await Form.doctor_name.set()
    await message.answer("Введите ФИО врача:")

@dp.message_handler(state=Form.doctor_name)
async def process_doctor_name(message: types.Message, state: FSMContext):
    await state.update_data(doctor_name=message.text)
    await Form.patient_name.set()
    await message.answer("Введите ФИО пациента:")

@dp.message_handler(state=Form.patient_name)
async def process_patient_name(message: types.Message, state: FSMContext):
    await state.update_data(patient_name=message.text)
    await Form.diagnosis.set()
    await message.answer("Введите диагноз:")

@dp.message_handler(state=Form.diagnosis)
async def process_diagnosis(message: types.Message, state: FSMContext):
    await state.update_data(diagnosis=message.text)
    await Form.lab_results.set()
    await message.answer("Введите лабораторные результаты:")

@dp.message_handler(state=Form.lab_results)
async def process_lab(message: types.Message, state: FSMContext):
    await state.update_data(lab_results=message.text)
    data = await state.get_data()

    # Генерация документа
    doc = Document()
    doc.add_heading("Экстренное извещение", level=1)
    doc.add_paragraph(f"Врач: {data['doctor_name']}")
    doc.add_paragraph(f"Пациент: {data['patient_name']}")
    doc.add_paragraph(f"Диагноз: {data['diagnosis']}")
    doc.add_paragraph(f"Лабораторные данные: {data['lab_results']}")
    
    filename = f"izvesh_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(filename)
    
    with open(filename, "rb") as file:
        await message.answer_document(file, caption="Вот ваше извещение 📄")

    os.remove(filename)

    await message.answer("Извещение сформировано. Чтобы отправить новое — нажмите \"📝 Подать извещение\".", 
                         reply_markup=start_keyboard)
    await state.finish()

@dp.message_handler()
async def fallback(message: types.Message):
    await message.answer("Чтобы подать извещение, нажмите \"📝 Подать извещение\".")

if __name__ == "__main__":
    executor.start_polling(dp, skip_updates=True)
